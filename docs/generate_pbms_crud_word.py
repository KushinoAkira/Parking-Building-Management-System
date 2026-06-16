from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def main():
    doc = Document()

    title = doc.add_heading("BASIC CRUD IN DB TABLES", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = doc.add_paragraph("Parking Building Management System (PBMS) - Master Data Setup")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_heading(doc, "1. Muc tieu bai tap", level=1)
    add_bullets(
        doc,
        [
            "Xac dinh va chuan hoa nhom master data cho PBMS.",
            "Thiet ke cau truc bang du lieu co kha nang mo rong, de bao tri.",
            "Cai dat day du luong Create - Read - Update - Delete cho tung bang.",
            "Khoi tao du lieu mau ban dau de he thong co the van hanh ngay.",
            "Ket noi CRUD voi API va giao dien quan tri.",
        ],
    )

    add_heading(doc, "2. Pham vi Master Data duoc chon cho PBMS", level=1)
    add_paragraph = doc.add_paragraph
    add_paragraph(
        "Trong PBMS, nhom de tai chon 4 bang master data de trien khai CRUD: "
        "VehicleTypes, ParkingZones, PricingPolicies, PaymentMethods (qua bang SystemConfigs)."
    )
    add_table(
        doc,
        ["Bang", "Muc dich", "Vi du du lieu"],
        [
            ["VehicleTypes", "Danh muc loai xe duoc chap nhan", "MOTORBIKE, CAR, EV"],
            ["ParkingZones", "Cau hinh khu vuc do xe theo loai xe", "A, B, C, B1"],
            ["PricingPolicies", "Quy tac tinh phi theo loai xe", "Gia gio, DailyMaxFee"],
            ["SystemConfigs", "Cau hinh phuong thuc thanh toan va bien he thong", "Cash, BankTransfer, EWallet"],
        ],
    )

    add_heading(doc, "3. Tieu chuan thiet ke database", level=1)
    add_bullets(
        doc,
        [
            "Moi bang co cac truong nen: id, name/code, description, status, created_at, updated_at (hoac cot tuong duong).",
            "Rang buoc unique cho cac truong nhan dien nghiep vu (TypeCode, ZoneCode).",
            "Status su dung gia tri co kiem soat: Active/Inactive/Maintenance/Locked.",
            "Dung chi muc cho cac cot tim kiem thuong xuyen de toi uu truy van.",
        ],
    )

    add_heading(doc, "4. Mau DDL de xay dung bang", level=1)
    doc.add_paragraph("4.1. VehicleTypes")
    doc.add_paragraph(
        "CREATE TABLE VehicleTypes (\n"
        "  VehicleTypeID INT PRIMARY KEY IDENTITY(1,1),\n"
        "  TypeCode VARCHAR(20) NOT NULL UNIQUE,\n"
        "  TypeName VARCHAR(100) NOT NULL,\n"
        "  Status VARCHAR(20) NOT NULL DEFAULT 'Active',\n"
        "  Description VARCHAR(255) NULL,\n"
        "  CreatedAt DATETIME NOT NULL DEFAULT GETDATE(),\n"
        "  UpdatedAt DATETIME NOT NULL DEFAULT GETDATE()\n"
        ");"
    )

    doc.add_paragraph("4.2. ParkingZones")
    doc.add_paragraph(
        "CREATE TABLE ParkingZones (\n"
        "  ZoneID INT PRIMARY KEY IDENTITY(1,1),\n"
        "  ZoneCode VARCHAR(20) NOT NULL UNIQUE,\n"
        "  ZoneName VARCHAR(100) NOT NULL,\n"
        "  VehicleTypeID INT NOT NULL FOREIGN KEY REFERENCES VehicleTypes(VehicleTypeID),\n"
        "  Capacity INT NOT NULL,\n"
        "  Status VARCHAR(20) NOT NULL DEFAULT 'Active',\n"
        "  CreatedAt DATETIME NOT NULL DEFAULT GETDATE(),\n"
        "  UpdatedAt DATETIME NOT NULL DEFAULT GETDATE()\n"
        ");"
    )

    add_heading(doc, "5. Yeu cau CRUD cho tung bang", level=1)
    add_heading(doc, "5.1 Create", level=2)
    add_bullets(
        doc,
        [
            "Bat buoc validate du lieu dau vao (khong rong, dung dinh dang, dung enum).",
            "Khong cho phep tao ban ghi bi trung code/ten nghiep vu.",
            "Tra ve thong bao loi ro rang neu vi pham rang buoc.",
        ],
    )
    add_heading(doc, "5.2 Read", level=2)
    add_bullets(
        doc,
        [
            "Ho tro danh sach + chi tiet theo id.",
            "Ho tro loc theo status va tim kiem theo keyword.",
            "Khuyen nghi phan trang voi du lieu lon.",
        ],
    )
    add_heading(doc, "5.3 Update", level=2)
    add_bullets(
        doc,
        [
            "Khong cho phep update thanh gia tri rong hoac trung voi ban ghi khac.",
            "Cap nhat cot UpdatedAt moi lan sua du lieu.",
            "Giu on dinh du lieu tham chieu (khong sua khoa chinh).",
        ],
    )
    add_heading(doc, "5.4 Delete", level=2)
    add_bullets(
        doc,
        [
            "Uu tien Soft Delete (chuyen Status = Inactive) thay vi xoa cung.",
            "Khong cho xoa neu da co du lieu phat sinh tham chieu neu he thong yeu cau bao toan.",
            "Ghi log thao tac xoa cho muc dich truy vet.",
        ],
    )

    add_heading(doc, "6. API de xuat (Backend ASP.NET Core)", level=1)
    add_table(
        doc,
        ["Bang", "Danh sach API REST"],
        [
            ["VehicleTypes", "GET /api/vehicle-types | GET /api/vehicle-types/{id} | POST | PUT | DELETE"],
            ["ParkingZones", "GET /api/parking-zones | GET /api/parking-zones/{id} | POST | PUT | DELETE"],
            ["PricingPolicies", "GET /api/pricing-policies | GET /api/pricing-policies/{id} | POST | PUT | DELETE"],
            ["SystemConfigs", "GET /api/system-configs | PUT /api/system-configs/{key}"],
        ],
    )

    add_heading(doc, "7. Du lieu mau khoi tao (toi thieu 5 dong/bang)", level=1)
    add_paragraph("7.1 VehicleTypes")
    add_table(
        doc,
        ["TypeCode", "TypeName", "Status"],
        [
            ["MOTORBIKE", "Xe may", "Active"],
            ["CAR", "O to", "Active"],
            ["EV", "Xe dien", "Active"],
            ["TRUCK", "Xe tai nhe", "Inactive"],
            ["VIP_CAR", "O to VIP", "Active"],
        ],
    )

    add_paragraph("7.2 Payment methods (SystemConfigs)")
    add_table(
        doc,
        ["ConfigKey", "ConfigValue", "Description"],
        [
            ["payment.method.1", "Cash", "Thanh toan tien mat"],
            ["payment.method.2", "BankTransfer", "Chuyen khoan ngan hang"],
            ["payment.method.3", "EWallet", "Thanh toan vi dien tu"],
            ["payment.method.4", "QRCode", "Thanh toan ma QR"],
            ["payment.method.5", "Card", "The ngan hang/credit"],
        ],
    )

    add_heading(doc, "8. Tieu chi nghiem thu", level=1)
    add_numbered(
        doc,
        [
            "Demo du 4 thao tac CRUD tren it nhat 3 bang master data.",
            "Validate dung cac rule bat buoc (required, unique, enum).",
            "Co du lieu seed ban dau va hien thi duoc tren giao dien/API.",
            "Khong phat sinh loi compile/build trong frontend va backend.",
            "Tai lieu nop day du gom mo ta, script SQL, source code, va minh chung demo.",
        ],
    )

    add_heading(doc, "9. San pham nop", level=1)
    add_numbered(
        doc,
        [
            "Tai lieu mo ta master data (file Word nay).",
            "File SQL gom CREATE TABLE + INSERT mau.",
            "Source code CRUD (API/UI).",
            "Anh/chup man hinh hoac video demo luong CRUD.",
            "Danh sach thanh vien nhom (Ho ten, MSSV, ty le dong gop).",
        ],
    )

    doc.add_paragraph("\nTai lieu duoc bien soan theo huong dan hoc phan SWP391 va can chinh sua nho de phu hop voi du an cua nhom truoc khi nop.")

    output_path = "docs/PBMS_Basic_CRUD_Master_Data_Setup.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
